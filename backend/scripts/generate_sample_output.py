"""
Standalone script: extract + generate quotation DOCX from an RFQ file.

Usage (from backend/):
    .venv/Scripts/python.exe scripts/generate_sample_output.py

Output: data/samples/output/EEPL-711-quotation.docx
"""

import os
import sys
from pathlib import Path

# ── Project root on sys.path ──────────────────────────────────────────────────
BACKEND_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BACKEND_DIR))

# Load .env before importing any app modules so Settings picks up values.
from dotenv import load_dotenv
load_dotenv(BACKEND_DIR / ".env")

# ── Imports ───────────────────────────────────────────────────────────────────
import json
from docx import Document as DocxDocument

from app.services.docx_exporter import create_quote_docx
from app.services.generator import generate_quote_from_structured
from app.services.extractor import extract_rfq_metadata

# ── Paths ─────────────────────────────────────────────────────────────────────
RFQ_PATH   = BACKEND_DIR / "data" / "samples" / "rfq" / "RFQ EEPL 711.docx"
OUTPUT_DIR = BACKEND_DIR / "data" / "samples" / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / "EEPL-711-quotation.docx"


def extract_text_from_docx(path: Path) -> str:
    """Simple DOCX text extractor using python-docx (no docling needed)."""
    doc = DocxDocument(str(path))
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    # Also pull text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def main():
    print(f"[1/4] Reading RFQ: {RFQ_PATH.name}")
    rfq_text = extract_text_from_docx(RFQ_PATH)
    print(f"      Extracted {len(rfq_text)} characters.")

    print("[2/4] Extracting metadata via LLM...")
    metadata = extract_rfq_metadata(rfq_text)
    print(f"      Metadata: {json.dumps(metadata, indent=2)}")

    # Build the rfq_data dict that the generator expects
    rfq_data = {
        **metadata,
        "raw_text": rfq_text,
        "source_file": RFQ_PATH.name,
    }

    print("[3/4] Generating quotation JSON via LLM...")
    quote_data = generate_quote_from_structured(rfq_data, context_docs=[])
    print(f"      Top-level keys: {list(quote_data.keys())}")

    # ── Normalise header ──────────────────────────────────────────────────────
    if not isinstance(quote_data.get("header"), dict):
        quote_data["header"] = {}
    hdr = quote_data["header"]
    hdr.setdefault("company_name", "ACME PROCESS SYSTEMS PVT LTD, PUNE (INDIA)")
    hdr.setdefault("company_address", "Gat No. 588, Near KSB, Urawade Road, Pirangut, Pune - 412115, Maharashtra (India)")
    hdr.setdefault("certifications", [
        "CERTIFIED COMPANY BY THE AMERICAN SOCIETY OF MECHANICAL ENGINEERS (ASME U STAMP)",
        "ISO 9001:2015 CERTIFIED COMPANY BY TUV NORD GERMANY",
    ])

    # ── Normalise offer_details ───────────────────────────────────────────────
    if not isinstance(quote_data.get("offer_details"), dict):
        quote_data["offer_details"] = {}
    offer = quote_data["offer_details"]
    offer.setdefault("customer_name", metadata.get("customer_name", ""))
    offer.setdefault("offer_date", "")
    offer.setdefault("prepared_by", "")

    # ── Normalise footer (LLM sometimes returns a plain string) ───────────────
    raw_footer = quote_data.get("footer", {})
    if not isinstance(raw_footer, dict):
        # Parse "For ACME...\n\n\nAuthorized Signatory" into structured dict
        footer_text = str(raw_footer)
        quote_data["footer"] = {
            "authorized_signatory": "Authorized Signatory",
            "title": "SR. ENGINEER - Application & Sales",
            "email": hdr.get("email", ""),
            "phone": hdr.get("phone", ""),
        }

    # ── Normalise covering_letter body (LLM sometimes returns a plain string) ─
    cl = quote_data.get("covering_letter", {})
    if isinstance(cl, dict) and isinstance(cl.get("body"), str):
        cl["body"] = [cl["body"]]

    # Save raw JSON for inspection
    json_path = OUTPUT_DIR / "EEPL-711-quotation-raw.json"
    json_path.write_text(json.dumps(quote_data, indent=2, default=str), encoding="utf-8")
    print(f"      Raw JSON saved to: {json_path.name}")

    print("[4/4] Building DOCX...")
    buf = create_quote_docx(quote_data)
    OUTPUT_PATH.write_bytes(buf.read())
    print(f"\n  Output: {OUTPUT_PATH}")
    print("  Done!")


if __name__ == "__main__":
    main()
