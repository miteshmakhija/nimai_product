import os
os.environ.setdefault("DATABASE_URL", "sqlite+pysqlite:///:memory:")

import pytest
import uuid
import io
from datetime import datetime
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool
from app.models.db import Base, DocxTemplateRecord


@pytest.fixture()
def db():
    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    Session = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    s = Session()
    try:
        yield s
    finally:
        s.close()


def test_docx_template_record_create(db):
    user_id = uuid.uuid4()
    row = DocxTemplateRecord(
        id=uuid.uuid4(),
        name="acme_v1.docx",
        template_blob=b"PK fake docx bytes",
        uploaded_by=user_id,
        uploaded_at=datetime.utcnow(),
        is_active=True,
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    assert row.name == "acme_v1.docx"
    assert row.is_active is True
    assert row.template_blob == b"PK fake docx bytes"


# ── Service tests ─────────────────────────────────────────────────────────────

from app.services.docx_template_service import get_active_template, save_template, deactivate_template


def test_save_template_sets_active(db):
    user_id = uuid.uuid4()
    row = save_template(db, blob=b"PK blob1", name="v1.docx", user_id=user_id)
    assert row.is_active is True
    assert row.name == "v1.docx"


def test_save_template_deactivates_previous(db):
    user_id = uuid.uuid4()
    first = save_template(db, blob=b"PK blob1", name="v1.docx", user_id=user_id)
    second = save_template(db, blob=b"PK blob2", name="v2.docx", user_id=user_id)
    db.refresh(first)
    assert first.is_active is False
    assert second.is_active is True


def test_get_active_template_returns_active(db):
    user_id = uuid.uuid4()
    save_template(db, blob=b"PK blob1", name="v1.docx", user_id=user_id)
    active = get_active_template(db)
    assert active is not None
    assert active.name == "v1.docx"


def test_get_active_template_returns_none_when_empty(db):
    assert get_active_template(db) is None


def test_deactivate_template(db):
    user_id = uuid.uuid4()
    save_template(db, blob=b"PK blob1", name="v1.docx", user_id=user_id)
    deactivate_template(db)
    assert get_active_template(db) is None


# ── Renderer tests ────────────────────────────────────────────────────────────

from app.services.docxtpl_renderer import render_template


def _make_minimal_docx() -> bytes:
    from docx import Document
    doc = Document()
    doc.add_paragraph("{{ header_company }}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_render_template_returns_bytes():
    blob = _make_minimal_docx()
    result = render_template(blob, {"header_company": "ACME Ltd"})
    assert isinstance(result, bytes)
    assert len(result) > 100


def test_render_template_bad_template_raises():
    with pytest.raises(Exception):
        render_template(b"not a docx", {})


# ── API router tests ──────────────────────────────────────────────────────────

from fastapi.testclient import TestClient
from app.api.main import app
from app.db.session import get_db
from app.models.db import User, UserRole
from app.core.security import create_access_token


def _make_admin(db_session) -> tuple:
    user = User(
        id=uuid.uuid4(),
        email=f"admin_{uuid.uuid4().hex[:6]}@test.com",
        full_name="Admin",
        role=UserRole.admin,
        is_active=True,
        auth_provider="local",
        password_hash="x",
    )
    db_session.add(user)
    db_session.commit()
    token = create_access_token({"sub": str(user.id)})
    return user, token


@pytest.fixture()
def api_client(db):
    app.dependency_overrides[get_db] = lambda: db
    with TestClient(app, raise_server_exceptions=False) as c:
        yield c
    app.dependency_overrides.clear()


def test_get_template_no_active(api_client, db):
    _, token = _make_admin(db)
    r = api_client.get("/docx-template", headers={"Authorization": f"Bearer {token}"})
    assert r.status_code == 200
    assert r.json()["active"] is False


def test_upload_and_get_template(api_client, db):
    from docx import Document
    doc = Document()
    doc.add_paragraph("{{ company }}")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    _, token = _make_admin(db)
    r = api_client.post(
        "/docx-template",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("test.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 200
    assert r.json()["name"] == "test.docx"

    r2 = api_client.get("/docx-template", headers={"Authorization": f"Bearer {token}"})
    assert r2.json()["active"] is True
    assert r2.json()["name"] == "test.docx"


def test_delete_template(api_client, db):
    _, token = _make_admin(db)
    r = api_client.delete("/docx-template", headers={"Authorization": f"Bearer {token}"})
    assert r.status_code == 200
