# app/services/docx_export.py
"""Generate Word (.docx) files from TipTap HTML or result_json."""
import io
import json
from html.parser import HTMLParser

from docx import Document
from docx.shared import Pt


class _HtmlToDocxParser(HTMLParser):
    """Minimal HTML → python-docx converter for TipTap output."""

    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self._para = None
        self._run = None
        self._bold = False
        self._italic = False
        self._list_type = None  # "ul" or "ol"
        self._list_counter = 0
        self._in_table = False
        self._table = None
        self._row = None
        self._cell = None
        self._skip_tags = {"html", "body", "head"}

    def handle_starttag(self, tag, attrs):
        if tag in self._skip_tags:
            return
        if tag in ("h1", "h2", "h3"):
            level = int(tag[1])
            self._para = self.doc.add_heading("", level=level)
        elif tag == "p":
            self._para = self.doc.add_paragraph()
        elif tag == "b" or tag == "strong":
            self._bold = True
        elif tag == "i" or tag == "em":
            self._italic = True
        elif tag == "ul":
            self._list_type = "ul"
        elif tag == "ol":
            self._list_type = "ol"
            self._list_counter = 0
        elif tag == "li":
            if self._list_type == "ol":
                self._list_counter += 1
                self._para = self.doc.add_paragraph(style="List Number")
            else:
                self._para = self.doc.add_paragraph(style="List Bullet")
        elif tag == "table":
            self._in_table = True
            self._table = None
        elif tag == "tr":
            if self._table is None:
                self._table = self.doc.add_table(rows=0, cols=1)
                self._table.style = "Table Grid"
            self._row = self._table.add_row()
            self._cell_index = 0
        elif tag in ("td", "th"):
            if self._row is not None:
                if self._cell_index >= len(self._row.cells):
                    self._cell = self._row.cells[-1]
                else:
                    self._cell = self._row.cells[self._cell_index]
                self._cell_index += 1
                self._para = self._cell.paragraphs[0]
                if tag == "th":
                    run = self._para.add_run("")
                    run.bold = True
        elif tag == "br":
            if self._para:
                self._para.add_run("\n")

    def handle_endtag(self, tag):
        if tag in ("b", "strong"):
            self._bold = False
        elif tag in ("i", "em"):
            self._italic = False
        elif tag in ("ul", "ol"):
            self._list_type = None
            self._list_counter = 0
        elif tag in ("td", "th"):
            self._para = None
        elif tag == "tr":
            self._row = None
        elif tag == "table":
            self._in_table = False

    def handle_data(self, data):
        text = data
        if not text:
            return
        if self._para is None:
            self._para = self.doc.add_paragraph()
        run = self._para.add_run(text)
        run.bold = self._bold
        run.italic = self._italic


def generate_docx_from_html(html: str, run) -> bytes:
    """Convert TipTap HTML to a .docx byte string."""
    doc = Document()
    # Add title from metadata if available
    title = " — ".join(filter(None, [
        getattr(run, "meta_company_name", None),
        getattr(run, "meta_product", None),
    ])) or f"Quote {str(run.id)[:8]}"
    doc.add_heading(title, level=0)

    parser = _HtmlToDocxParser(doc)
    parser.feed(html)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_docx_from_result_json(result_json: dict, run) -> bytes:
    """Convert structured result_json to a .docx byte string (fallback)."""
    doc = Document()
    title = " — ".join(filter(None, [
        getattr(run, "meta_company_name", None),
        getattr(run, "meta_product", None),
    ])) or f"Quote {str(run.id)[:8]}"
    doc.add_heading(title, level=0)

    if not result_json:
        doc.add_paragraph("No content generated.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _add_section(heading: str, data):
        doc.add_heading(heading, level=1)
        if isinstance(data, str):
            doc.add_paragraph(data)
        elif isinstance(data, list):
            for item in data:
                if isinstance(item, dict):
                    doc.add_paragraph(
                        " | ".join(f"{k}: {v}" for k, v in item.items()),
                        style="List Bullet",
                    )
                else:
                    doc.add_paragraph(str(item), style="List Bullet")
        elif isinstance(data, dict):
            for k, v in data.items():
                doc.add_paragraph(f"{k}: {v}")
        else:
            doc.add_paragraph(str(data))

    for section_key, section_val in result_json.items():
        heading = section_key.replace("_", " ").title()
        _add_section(heading, section_val)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
