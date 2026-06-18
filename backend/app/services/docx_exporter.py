import ast
import io
import json
import re
from collections.abc import Mapping
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[2]
DEFAULT_LOGO_PATH = BASE_DIR / "data" / "assets" / "acme_logo.png"
DEFAULT_WATERMARK_PATH = BASE_DIR / "data" / "assets" / "acme_watermark.png"

FONT_NAME = "Times New Roman"
HEADER_SHADE = "D9E2F3"
LABEL_SHADE = "BFBFBF"


def create_quote_docx(quote_data: dict) -> io.BytesIO:
    """
    Convert structured quote data to a DOCX that follows the ACME sample quotation
    layout: repeated document-control header, cover page, bordered spec tables,
    general notes, price summary, commercial terms, and signature block.
    """
    doc = Document()

    header_info = quote_data.get("header", {})
    offer = quote_data.get("offer_details", {})
    equipment = quote_data.get("equipment", [])
    general_notes = quote_data.get("general_notes", [])
    pricing = quote_data.get("pricing_table", {})
    terms = quote_data.get("terms_and_conditions", [])
    footer = quote_data.get("footer", {})
    covering_letter = quote_data.get("covering_letter", {})
    revision_history = quote_data.get("revision_history", [])

    _configure_document(doc)
    _configure_repeating_header(doc, header_info, offer)

    _build_title_page(doc, header_info, offer)

    if not _is_empty_value(covering_letter):
        doc.add_page_break()
        _build_covering_letter_page(doc, offer, covering_letter, footer)

    for item in equipment:
        if _is_empty_value(item):
            continue
        doc.add_page_break()
        _build_technical_page(doc, item, revision_history)

    if not _is_empty_value(general_notes):
        doc.add_page_break()
        _build_general_notes_page(doc, general_notes)

    doc.add_page_break()
    _build_pricing_page(doc, header_info, pricing, terms, footer)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(5.0)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.55)

    _set_page_border(section)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.5)


def _configure_repeating_header(doc, header_info: dict, offer: dict):
    header = doc.sections[0].header
    for paragraph in header.paragraphs:
        paragraph.text = ""

    table = header.add_table(rows=4, cols=5, width=Cm(18.2))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    logo_cell = table.cell(0, 0)
    logo_cell.merge(table.cell(3, 0))
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    company_cell = table.cell(0, 1)
    company_cell.merge(table.cell(0, 4))

    title_cell = table.cell(1, 1)
    title_cell.merge(table.cell(3, 1))
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    logo_path = _resolve_image_path(header_info, "logo_path", DEFAULT_LOGO_PATH)
    if logo_path:
        paragraph = logo_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(logo_path), width=Inches(1.55))
    else:
        _set_cell_text(logo_cell, "ACME\nPROCESS : PEOPLE : PLANET", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)

    company = header_info.get("company_name", "ACME PROCESS SYSTEMS PVT LTD, PUNE (INDIA)")
    _set_cell_text(company_cell, _spaced_text(company), size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    title = f"TECHNO-\nCOMMERCIAL OFFER\nFOR\n{_equipment_title_from_offer(offer).upper()}"
    _set_cell_text(title_cell, title, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade_cell(title_cell, LABEL_SHADE)

    rows = [
        ("ISSUE:", offer.get("issue_or_version") or offer.get("version") or "Version : 1"),
        ("DOCUMENT No.", offer.get("document_no") or offer.get("offer_number", "")),
        ("DATE :", offer.get("offer_date", "")),
        ("PAGE No.", None),
    ]
    for row_index, (label, value) in enumerate(rows):
        _set_cell_text(table.cell(row_index, 2), label, size=9)
        _shade_cell(table.cell(row_index, 2), LABEL_SHADE)
        value_cell = table.cell(row_index, 3)
        value_cell.merge(table.cell(row_index, 4))
        if value is None:
            _set_page_number_cell(value_cell)
        else:
            _set_cell_text(value_cell, _stringify_value(value), size=10)

    for row in table.rows:
        row.height = Cm(0.72)


def _build_title_page(doc, header_info: dict, offer: dict):
    address = header_info.get("company_address", "")
    certifications = header_info.get(
        "certifications",
        [
            "CERTIFIED COMPANY BY THE AMERICAN SOCIETY OF MECHANICAL ENGINEERS (ASME U STAMP)",
            "ISO 9001:2015 CERTIFIED COMPANY BY TUV NORD GERMANY",
        ],
    )

    for cert in certifications:
        if _is_empty_value(cert):
            continue
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(str(cert).upper())
        _style_run(run, bold=True, underline=True, size=13)
        if "ASME" in str(cert).upper() or "ISO" in str(cert).upper():
            _add_spacer(doc, 8)

    _add_centered_image(doc, _resolve_image_path(header_info, "watermark_path", DEFAULT_WATERMARK_PATH), Inches(2.5))

    if not _is_empty_value(address):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(_stringify_value(address))
        _style_run(run, size=13)

    doc.add_paragraph()

    offer_box = doc.add_table(rows=1, cols=1)
    offer_box.style = "Table Grid"
    offer_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(offer_box, Cm(13.5))
    _set_cell_text(
        offer_box.cell(0, 0),
        f"TECHNO-COMMERCIAL OFFER FOR\n{_equipment_title_from_offer(offer).upper()}",
        bold=True,
        size=15,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.add_paragraph()
    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(table, Cm(18.2))

    rows = [
        ("OFFER NO.", offer.get("offer_number", ""), "OFFER DATE :", offer.get("offer_date", "")),
        ("PREPARED BY", offer.get("prepared_by", ""), "CONTACT PERSON:", offer.get("contact_person", "")),
        ("CLIENT :", offer.get("customer_name", ""), "", ""),
    ]
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        _set_cell_text(row.cells[0], values[0], bold=True, size=9)
        _set_cell_text(row.cells[1], _stringify_value(values[1]), size=9)
        _set_cell_text(row.cells[2], values[2], bold=True, size=9)
        _set_cell_text(row.cells[3], _stringify_value(values[3]), size=9)
        if not values[2]:
            row.cells[1].merge(row.cells[3])


def _build_covering_letter_page(doc, offer: dict, letter: dict, footer: dict):
    to_value = letter.get("to") or offer.get("customer_name", "")
    if not _is_empty_value(to_value):
        doc.add_paragraph("To,")
        doc.add_paragraph(_stringify_value(to_value))

    attention = letter.get("kind_attention") or offer.get("contact_person", "")
    if not _is_empty_value(attention):
        para = doc.add_paragraph()
        _add_labeled_text(para, "Kind Attention:", attention)

    subject = letter.get("subject") or offer.get("subject", "")
    if not _is_empty_value(subject):
        para = doc.add_paragraph()
        _add_labeled_text(para, "Subject:", subject)

    doc.add_paragraph()
    doc.add_paragraph(letter.get("salutation", "Dear Sir,"))

    body = letter.get("body") or [
        "We thank you for your enquiry for subject equipment. Based on the data provided by you, we are pleased to submit our Techno Commercial offer for the same as below -",
        "Kindly find attached herewith detailed Techno-Commercial Proposal in line with enquiry documents provided to us and discussions with our technical team.",
        "We hope that the offer is in line with your requirements and we look forward to receiving your valuable order.",
        "Please feel free to contact Undersign and our team for any queries if any.",
    ]
    for line in _as_list(body):
        if _is_empty_value(line):
            continue
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(1.1)
        _style_run(para.add_run(str(line)), size=10.5)

    doc.add_paragraph()
    doc.add_paragraph("Best Regards,")
    doc.add_paragraph("FOR ACME PROCESS SYSTEMS PVT. LTD.")

    signatory = footer.get("authorized_signatory") or offer.get("prepared_by", "")
    if not _is_empty_value(signatory):
        _style_run(doc.add_paragraph().add_run(str(signatory)), bold=True, size=10.5)
    if not _is_empty_value(footer.get("title")):
        doc.add_paragraph(str(footer.get("title")))
    if not _is_empty_value(footer.get("phone")):
        doc.add_paragraph(f"Contact- {footer.get('phone')}")
    if not _is_empty_value(footer.get("email")):
        doc.add_paragraph(f"Email- {footer.get('email')}")


def _build_technical_page(doc, item: dict, revision_history=None):
    name = item.get("name", "Equipment")
    moc = item.get("moc", "")
    tag = item.get("tag_no", "")

    title = str(name)
    if not _is_empty_value(moc) and str(moc).lower() not in title.lower():
        title = f"{title} ({moc})"

    para = doc.add_paragraph()
    run = para.add_run(title)
    _style_run(run, bold=True, size=12)

    spec_table = doc.add_table(rows=0, cols=2)
    spec_table.style = "Table Grid"
    spec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(spec_table, Cm(17.7))

    if not _is_empty_value(tag):
        row = spec_table.add_row()
        row.cells[0].merge(row.cells[1])
        _set_cell_text(row.cells[0], f"Tag No : {_stringify_value(tag)}", bold=True, size=9)
        _shade_cell(row.cells[0], HEADER_SHADE)

    sections = [
        ("Process & Mechanical Details", item.get("process_and_mechanical_details", {}), "Vessel"),
        ("Vessel Details :", item.get("vessel_details", {}), None),
        ("Exchanger Details", item.get("exchanger_details", {}), None),
        ("Material of Construction", item.get("material_of_construction", {}), None),
        ("Limpet / Jacket Details", item.get("limpet_or_jacket_details", {}), None),
        ("Supports", item.get("supports", {}), None),
        ("Accessories", item.get("accessories", {}), None),
        ("Nozzle Schedule", item.get("nozzle_schedule", {}), None),
        ("AGITATOR DETAILS", item.get("agitator_details", {}), None),
        ("Instrumentation", item.get("instrumentation", {}), None),
        ("Surface Finish", item.get("surface_finish", {}), None),
    ]

    pending_record_tables = []
    for heading, data, value_heading in sections:
        data = _normalize_structured_value(data)
        if _is_empty_value(data):
            continue
        if _is_record_list(data):
            pending_record_tables.append((heading, data))
        else:
            _append_key_values_to_table(spec_table, heading, data, value_heading)

    qty = item.get("quantity", "")
    if not _is_empty_value(qty):
        row = spec_table.add_row()
        _set_cell_text(row.cells[0], "Quantity", size=9)
        _set_cell_text(row.cells[1], _stringify_value(qty), size=9)

    for heading, rows in pending_record_tables:
        doc.add_paragraph()
        _add_section_heading(doc, heading)
        if heading == "Nozzle Schedule":
            _add_nozzle_schedule_table(doc, rows)
        else:
            _add_record_table(doc, rows)

    item_revision_history = item.get("revision_history") or revision_history
    if not _is_empty_value(item_revision_history):
        doc.add_paragraph()
        _add_section_heading(doc, "Revision History")
        _add_record_table(doc, item_revision_history)


def _build_general_notes_page(doc, notes: list):
    heading_para = doc.add_paragraph()
    heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_run(heading_para.add_run("GENERAL NOTES"), bold=True, underline=True, size=11)

    for section in _as_list(notes):
        if _is_empty_value(section):
            continue
        if isinstance(section, Mapping):
            section_title = section.get("heading", "")
            if not _is_empty_value(section_title):
                _style_run(doc.add_paragraph().add_run(str(section_title)), bold=True, size=11)
            items = section.get("items", [])
        else:
            items = [section]

        for note in _as_list(items):
            if _is_empty_value(note):
                continue
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Cm(0.55)
            _style_run(para.add_run(_stringify_value(note)), size=10)


def _build_pricing_page(doc, header_info: dict, pricing: dict, terms: list, footer: dict):
    doc.add_paragraph()
    heading_para = doc.add_paragraph()
    heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_run(heading_para.add_run("PRICE SUMMARY"), bold=True, underline=True, size=11)

    items = pricing.get("items", []) if isinstance(pricing, Mapping) else []
    if items:
        _add_price_table(doc, pricing, items)

    if terms:
        doc.add_paragraph()
        _style_run(doc.add_paragraph().add_run("Commercial Terms"), bold=True, underline=True, size=11)
        for term in terms:
            if _is_empty_value(term):
                continue
            if isinstance(term, Mapping):
                _add_term_paragraph(doc, term.get("label", ""), term.get("value", ""))
            elif ":" in str(term):
                label, value = str(term).split(":", 1)
                _add_term_paragraph(doc, label.strip(), value.strip())
            else:
                _style_run(doc.add_paragraph().add_run(str(term)), size=10.5)

    doc.add_paragraph()
    closing = doc.add_paragraph()
    _style_run(
        closing.add_run(
            "We trust this is in line with your requirements. Should you require any further information then please feel free to contact us."
        ),
        size=10.5,
    )

    doc.add_paragraph()
    _style_run(doc.add_paragraph().add_run("Thanks & regards,"), size=10.5)
    doc.add_paragraph()

    signatory = footer.get("authorized_signatory", "")
    if not _is_empty_value(signatory):
        _style_run(doc.add_paragraph().add_run(str(signatory).upper()), bold=True, size=10.5)

    title = footer.get("title", "SR. ENGINEER - Application & Sales")
    if not _is_empty_value(title):
        _style_run(doc.add_paragraph().add_run(str(title).upper()), bold=True, size=10.5)

    email = footer.get("email", header_info.get("email", ""))
    if not _is_empty_value(email):
        para = doc.add_paragraph()
        _style_run(para.add_run("E- Mail: "), bold=True, size=10.5)
        _style_run(para.add_run(str(email)), bold=True, underline=True, size=10.5, color=RGBColor(0, 0, 255))

    phone = footer.get("phone", header_info.get("phone", ""))
    if not _is_empty_value(phone):
        _style_run(doc.add_paragraph().add_run(f"M.no: {phone}"), bold=True, size=10.5)

    company = header_info.get("company_name", "ACME PROCESS SYSTEM PVT. LTD")
    company = str(company).replace("SYSTEMS", "SYSTEM")
    _style_run(doc.add_paragraph().add_run(company.upper()), bold=True, size=10.5)


def _append_key_values_to_table(table, heading: str, data, value_heading=None):
    data = _normalize_structured_value(data)
    if _is_empty_value(data):
        return

    rows = []
    if isinstance(data, Mapping):
        rows = [(key, value) for key, value in data.items() if not _is_empty_value(value)]
    elif isinstance(data, list):
        rows = [(f"Item {index + 1}", value) for index, value in enumerate(data) if not _is_empty_value(value)]
    else:
        rows = [(heading, data)]

    if not rows:
        return

    header_row = table.add_row()
    _set_cell_text(header_row.cells[0], heading, bold=True, size=9)
    if value_heading:
        _set_cell_text(header_row.cells[1], value_heading, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        _set_cell_text(header_row.cells[1], "", size=9)

    for key, value in rows:
        row = table.add_row()
        _set_cell_text(row.cells[0], _humanize_key(key), size=9)
        _set_cell_text(row.cells[1], _stringify_value(value), size=9)


def _add_price_table(doc, pricing: dict, items: list):
    has_capacity = any(
        isinstance(item, Mapping) and not _is_empty_value(item.get("capacity") or item.get("capacity_or_size"))
        for item in items
    )
    columns = ["SR.\nNO", "ITEMS"]
    if has_capacity:
        columns.append("Capacity")
    columns.extend(["MOC", "PRICE EACH", "QTY", "TOTAL PRICE"])

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(table, Cm(17.5))

    for idx, column in enumerate(columns):
        _set_cell_text(table.rows[0].cells[idx], column, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(table.rows[0].cells[idx], HEADER_SHADE)

    for item in items:
        if not isinstance(item, Mapping):
            item = {"description": item}
        values = [item.get("sr_no", ""), item.get("description", "")]
        if has_capacity:
            values.append(item.get("capacity") or item.get("capacity_or_size", ""))
        values.extend([item.get("moc", ""), item.get("unit_price", ""), item.get("quantity", ""), item.get("total_price", "")])
        row = table.add_row()
        for idx, value in enumerate(values):
            _set_cell_text(row.cells[idx], _stringify_value(value), bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    total = pricing.get("subtotal", pricing.get("total", ""))
    if not _is_empty_value(total):
        row = table.add_row()
        row.cells[0].merge(row.cells[-2])
        _set_cell_text(row.cells[0], "TOTAL (FOR)", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(row.cells[-1], _stringify_value(total), bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    in_words = pricing.get("total_in_words", "")
    if not _is_empty_value(in_words):
        row = table.add_row()
        row.cells[0].merge(row.cells[-1])
        _set_cell_text(row.cells[0], f"IN WORDS: {_stringify_value(in_words)}", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_nozzle_schedule_table(doc, rows: list):
    columns = [
        ("nozzle", "Nozzle"),
        ("description", "Description"),
        ("size_nb", "Size, NB"),
        ("rating", "Rating"),
        ("schedule_or_thickness", "Sch./Thk"),
        ("quantity", "Qty"),
        ("remarks", "Remark"),
    ]
    _add_record_table(doc, rows, columns=columns)


def _add_record_table(doc, rows: list, columns=None):
    rows = _normalize_structured_value(rows)
    clean_rows = [row for row in _as_list(rows) if not _is_empty_value(row)]
    if not clean_rows:
        return

    if not all(isinstance(row, Mapping) for row in clean_rows):
        _append_key_values_to_table(doc.add_table(rows=0, cols=2), "Items", clean_rows)
        return

    if columns is None:
        seen = []
        for row in clean_rows:
            for key in row:
                if key not in seen and any(not _is_empty_value(other.get(key)) for other in clean_rows):
                    seen.append(key)
        columns = [(key, _humanize_key(key)) for key in seen]

    active_columns = [
        (key, label)
        for key, label in columns
        if any(not _is_empty_value(row.get(key)) for row in clean_rows)
    ]
    if not active_columns:
        return

    table = doc.add_table(rows=1, cols=len(active_columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(table, Cm(17.5))

    for idx, (_, label) in enumerate(active_columns):
        _set_cell_text(table.rows[0].cells[idx], label, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(table.rows[0].cells[idx], HEADER_SHADE)

    for row_data in clean_rows:
        row = table.add_row()
        for idx, (key, _) in enumerate(active_columns):
            _set_cell_text(row.cells[idx], _stringify_value(row_data.get(key, "")), size=8)


def _add_section_heading(doc, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    _style_run(para.add_run(text), bold=True, size=10)


def _add_term_paragraph(doc, label, value):
    if _is_empty_value(label) and _is_empty_value(value):
        return
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.65)
    if not _is_empty_value(label):
        _style_run(para.add_run(_stringify_value(label)), size=10.5)
    if not _is_empty_value(value):
        _style_run(para.add_run(f" : {_stringify_value(value)}"), size=10.5)


def _set_cell_text(cell, text, bold=False, underline=False, size=10, align=None):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    run = para.add_run(_stringify_value(text))
    _style_run(run, bold=bold, underline=underline, size=size)


def _style_run(run, bold=False, underline=False, size=10.5, color=None):
    run.bold = bold
    run.underline = underline
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _shade_cell(cell, color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._element.get_or_add_tcPr().append(shading)


def _set_table_width(table, width):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(int(width.pt * 20)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)


def _set_page_border(section):
    sect_pr = section._sectPr
    page_borders = sect_pr.find(qn("w:pgBorders"))
    if page_borders is None:
        page_borders = OxmlElement("w:pgBorders")
        sect_pr.append(page_borders)
    page_borders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        border = page_borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            page_borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "24")
        border.set(qn("w:color"), "000000")


def _set_page_number_cell(cell):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _style_run(para.add_run("Page "), size=10)
    _add_field(para, "PAGE")
    _style_run(para.add_run(" of "), size=10)
    _add_field(para, "NUMPAGES")


def _add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    _style_run(run, size=10)


def _add_centered_image(doc, path, width):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path:
        para.add_run().add_picture(str(path), width=width)


def _add_spacer(doc, points):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(points)


def _add_labeled_text(para, label, value):
    _style_run(para.add_run(label), bold=True, size=10.5)
    _style_run(para.add_run(f" {_stringify_value(value)}"), size=10.5)


def _resolve_image_path(header_info, key, default_path):
    value = header_info.get(key) if isinstance(header_info, Mapping) else None
    path = Path(value) if value else default_path
    if not path.is_absolute():
        path = BASE_DIR / path
    return path if path.exists() else None


def _spaced_text(text):
    text = str(text or "").upper().strip()
    text = re.sub(r"\s+", " ", text)
    return " ".join(text)


def _equipment_title_from_offer(offer: dict) -> str:
    subject = str(offer.get("subject") or "EQUIPMENT").strip()
    subject = re.sub(r"^TECHNO\s*-?\s*COMMERCIAL\s+OFFER\s+FOR\s+", "", subject, flags=re.I)
    return subject.strip() or "EQUIPMENT"


def _is_record_list(value) -> bool:
    value = _normalize_structured_value(value)
    return isinstance(value, list) and bool(value) and all(isinstance(row, Mapping) for row in value)


def _is_empty_value(value) -> bool:
    if value is None:
        return True
    if isinstance(value, str):
        cleaned = value.strip().lower()
        return cleaned in {"", "na", "n/a", "none", "null", "-", "not applicable"}
    if isinstance(value, Mapping):
        return all(_is_empty_value(v) for v in value.values())
    if isinstance(value, (list, tuple, set)):
        return all(_is_empty_value(v) for v in value)
    return False


def _as_list(value):
    value = _normalize_structured_value(value)
    if value is None:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, tuple):
        return list(value)
    return [value]


def _stringify_value(value) -> str:
    value = _normalize_structured_value(value)
    if value is None:
        return ""
    if isinstance(value, Mapping):
        parts = []
        for key, val in value.items():
            if not _is_empty_value(val):
                parts.append(f"{_humanize_key(key)}: {_stringify_value(val)}")
        return "; ".join(parts)
    if isinstance(value, list):
        return "; ".join(_stringify_value(v) for v in value if not _is_empty_value(v))
    return str(value)


def _humanize_key(key) -> str:
    text = str(key).replace("_", " ").strip()
    return " ".join(word.upper() if word.lower() in {"nb", "moc", "qty"} else word.capitalize() for word in text.split())


def _normalize_structured_value(value):
    """Convert legacy stringified dict/list payloads into structured values."""
    if isinstance(value, str):
        text = value.strip()
        if (text.startswith("{") and text.endswith("}")) or (text.startswith("[") and text.endswith("]")):
            for parser in (ast.literal_eval, json.loads):
                try:
                    parsed = parser(text)
                except (ValueError, SyntaxError, TypeError, json.JSONDecodeError):
                    continue
                return _normalize_structured_value(parsed)
        return value

    if isinstance(value, Mapping):
        return {key: _normalize_structured_value(val) for key, val in value.items()}

    if isinstance(value, list):
        return [_normalize_structured_value(item) for item in value]

    if isinstance(value, tuple):
        return [_normalize_structured_value(item) for item in value]

    return value
